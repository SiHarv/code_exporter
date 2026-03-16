"""DOCX exporter implementation."""

from __future__ import annotations

from typing import Callable

from docx import Document
from docx.shared import Pt

from scanner.file_scanner import SourceFile


def export_to_docx(
    output_path: str,
    source_files: list[SourceFile],
    summary_section: str,
    progress_callback: Callable[[int, int], None] | None = None,
) -> None:
    """Write exported code to a DOCX document with beginner-friendly formatting."""
    document = Document()
    total = len(source_files)

    document.add_heading("Code Export", level=0)
    intro = document.add_paragraph("Source files collected from the selected root folder.")
    intro.runs[0].font.size = Pt(11)

    for index, src_file in enumerate(source_files, start=1):
        # User requested file names as bold text, size 12.
        file_label = document.add_paragraph()
        file_run = file_label.add_run(src_file.relative_path)
        file_run.bold = True
        file_run.font.size = Pt(12)

        code_paragraph = document.add_paragraph(src_file.content)
        for run in code_paragraph.runs:
            run.font.name = "Consolas"
            run.font.size = Pt(9)

        document.add_paragraph("")

        if progress_callback:
            progress_callback(index, total)

    document.add_heading("Export Summary", level=1)
    for line in summary_section.splitlines():
        if not line.strip():
            continue
        summary_paragraph = document.add_paragraph(line)
        for run in summary_paragraph.runs:
            run.font.size = Pt(10)

    document.save(output_path)
